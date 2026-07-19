#!/usr/bin/env python3
"""将产品说明手册 Markdown + 截图生成为精美排版 Word 文档。"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Cm, Inches, Pt, RGBColor, Twips, Emu
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "outputs" / "产品说明手册.md"
SHOT_DIR = ROOT / "outputs" / "manual-screenshots"
OUT_PATH = ROOT / "outputs" / "猎头招聘交付系统-产品说明手册.docx"

# 产品冰蓝主题
BLUE = RGBColor(0x0B, 0x4F, 0x8A)
BLUE_ACCENT = RGBColor(0x02, 0x84, 0xC7)
DARK = RGBColor(0x1E, 0x29, 0x3B)
MUTED = RGBColor(0x47, 0x55, 0x69)
LIGHT_BLUE_HEX = "E8F4FF"
HEADER_BLUE_HEX = "0B4F8A"
TABLE_ALT_HEX = "F0F9FF"


def set_run_font(run, name="微软雅黑", size=11, bold=False, color=None, east_asia=None):
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), east_asia or name)


def set_paragraph_spacing(p, before=0, after=6, line=1.15):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), HEADER_BLUE_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color="CBD5E1"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def configure_styles(doc: Document):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "微软雅黑"
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color, before, after in [
        ("Heading 1", 18, BLUE, 18, 10),
        ("Heading 2", 14, BLUE_ACCENT, 14, 8),
        ("Heading 3", 12, DARK, 10, 6),
    ]:
        st = styles[style_name]
        st.font.name = "微软雅黑"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.15


def setup_page(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # 页眉
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("猎头招聘交付系统 · 产品说明手册")
    set_run_font(run, size=9, color=MUTED)
    add_horizontal_line(hp)

    # 页脚页码
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("— ")
    set_run_font(run, size=9, color=MUTED)
    # PAGE field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run2 = fp.add_run()
    run2._r.append(fldChar1)
    run2._r.append(instr)
    run2._r.append(fldChar2)
    set_run_font(run2, size=9, color=MUTED)
    run3 = fp.add_run(" —")
    set_run_font(run3, size=9, color=MUTED)


def add_cover(doc: Document):
    for _ in range(3):
        doc.add_paragraph()

    # 顶部色条表
    bar = doc.add_table(rows=1, cols=1)
    bar.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = bar.cell(0, 0)
    shade_cell(cell, HEADER_BLUE_HEX)
    cell.width = Cm(16.5)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=8, after=8)
    run = p.add_run("AI 招聘管理平台")
    set_run_font(run, size=12, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, before=24, after=8)
    run = title.add_run("产品说明手册")
    set_run_font(run, size=32, bold=True, color=BLUE)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(sub, before=0, after=18)
    run = sub.add_run("猎头招聘交付系统 · 岗位全生命周期管理")
    set_run_font(run, size=14, color=BLUE_ACCENT)

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_horizontal_line(line)

    meta_items = [
        ("文档性质", "产品功能说明（What / Why）"),
        ("适用对象", "产品、实施、培训、客户验收、新成员上手"),
        ("版本日期", "2026-07-19"),
        ("截图说明", "各功能模块附系统实机截图"),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=4, after=4)
        r1 = p.add_run(f"{label}  ·  ")
        set_run_font(r1, size=11, color=MUTED)
        r2 = p.add_run(value)
        set_run_font(r2, size=11, color=DARK)

    for _ in range(6):
        doc.add_paragraph()

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = foot.add_run("内部资料 · 请勿外传")
    set_run_font(run, size=10, color=MUTED)

    doc.add_page_break()


def add_toc_page(doc: Document):
    h = doc.add_heading("目录", level=1)
    for run in h.runs:
        set_run_font(run, size=18, bold=True, color=BLUE)

    note = doc.add_paragraph()
    set_paragraph_spacing(note, after=12)
    run = note.add_run("提示：在 Word 中右键目录区域 →「更新域」可刷新页码。")
    set_run_font(run, size=9, color=MUTED)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-2" \\h \\z \\u '
    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(qn("w:fldCharType"), "separate")
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_separate)
    # placeholder text until update
    t = OxmlElement("w:t")
    t.text = "（打开文档后请更新目录域）"
    run._r.append(t)
    run._r.append(fldChar_end)

    doc.add_page_break()


def add_body_paragraph(doc: Document, text: str, *, bold_segments=True):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=8, line=1.25)
    # simple **bold** support
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if bold_segments and part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run_font(run, size=11, bold=True, color=DARK)
        else:
            run = p.add_run(part)
            set_run_font(run, size=11, color=DARK)
    return p


def add_bullet(doc: Document, text: str, numbered=False, index=1):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=4, line=1.2)
    p.paragraph_format.left_indent = Cm(0.5)
    prefix = f"{index}. " if numbered else "• "
    run = p.add_run(prefix)
    set_run_font(run, size=11, color=BLUE_ACCENT, bold=True)
    # remaining with bold support
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run_font(run, size=11, bold=True, color=DARK)
        else:
            run = p.add_run(part)
            set_run_font(run, size=11, color=DARK)
    return p


def add_code_block(doc: Document, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, "F8FAFC")
    set_cell_borders(cell, "E2E8F0")
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=6, after=6, line=1.2)
    run = p.add_run(text.strip())
    set_run_font(run, name="Consolas", size=9, color=DARK, east_asia="微软雅黑")
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    if not headers:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, HEADER_BLUE_HEX)
        set_cell_borders(cell, HEADER_BLUE_HEX)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=4, after=4)
        run = p.add_run(h.strip())
        set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    for r_idx, row in enumerate(rows):
        for c_idx in range(len(headers)):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                shade_cell(cell, TABLE_ALT_HEX)
            set_cell_borders(cell, "CBD5E1")
            val = row[c_idx].strip() if c_idx < len(row) else ""
            val = val.replace("✅", "是").replace("❌", "否")
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before=3, after=3)
            run = p.add_run(val)
            set_run_font(run, size=9.5, color=DARK)

    doc.add_paragraph()


def add_image(doc: Document, path: Path, caption: str | None = None):
    if not path.exists():
        p = doc.add_paragraph()
        run = p.add_run(f"【截图缺失：{path.name}】")
        set_run_font(run, size=10, color=RGBColor(0xDC, 0x26, 0x26))
        return

    # 外框表
    frame = doc.add_table(rows=1, cols=1)
    frame.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = frame.cell(0, 0)
    set_cell_borders(cell, "BFDBFE")
    shade_cell(cell, "FFFFFF")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=4, after=4)
    run = p.add_run()
    # 最大宽度约 16cm
    run.add_picture(str(path), width=Cm(15.8))

    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(cap, before=4, after=12)
        run = cap.add_run(f"图  {caption}")
        set_run_font(run, size=9, color=MUTED)


def should_skip_section_line(line: str) -> bool:
    """跳过截图占位元信息，保留正文。"""
    skip_prefixes = (
        "#### 📷 截图占位",
        "| **建议文件**",
        "| **拍摄说明**",
        "| **可选补拍**",
        "| **建议补拍**",
        "| **插入方式**",
        "| 项 | 说明 |",
        "|----|------|",
        "| --- | --- |",
    )
    s = line.strip()
    if s.startswith(skip_prefixes):
        return True
    if s.startswith("| **建议") or s.startswith("| **拍摄") or s.startswith("| **可选") or s.startswith("| **插入"):
        return True
    if "建议文件" in s and s.startswith("|"):
        return True
    if "拍摄说明" in s and s.startswith("|"):
        return True
    if s.startswith("### 截图填写约定"):
        return True
    return False


def parse_table_block(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int]:
    headers = [c.strip() for c in lines[start].strip().strip("|").split("|")]
    i = start + 1
    if i < len(lines) and re.match(r"^\|?\s*:?-+", lines[i].strip()):
        i += 1
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        cols = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        # skip meta screenshot tables
        joined = " ".join(cols)
        if any(k in joined for k in ("建议文件", "拍摄说明", "可选补拍", "建议补拍", "插入方式")):
            i += 1
            continue
        if cols == headers or (len(cols) >= 2 and cols[0] == "项" and cols[1] == "说明"):
            # keep 用途/入口 tables
            pass
        rows.append(cols)
        i += 1
    return headers, rows, i


def is_meta_table(headers: list[str], rows: list[list[str]]) -> bool:
    blob = " ".join(headers) + " " + " ".join(" ".join(r) for r in rows)
    if "建议文件" in blob or "拍摄说明" in blob:
        return True
    return False


def convert_md(doc: Document, md_text: str):
    lines = md_text.splitlines()
    i = 0
    # skip title + blockquote + toc until first real ## 1.
    while i < len(lines):
        if lines[i].startswith("## 1."):
            break
        i += 1

    in_code = False
    code_buf: list[str] = []
    skip_until_heading = False
    fig_no = 1

    while i < len(lines):
        line = lines[i]
        raw = line.rstrip()
        s = raw.strip()

        if s.startswith("### 截图填写约定"):
            # skip until next ## or ### that is not 截图
            i += 1
            while i < len(lines):
                t = lines[i].strip()
                if t.startswith("## ") or (t.startswith("### ") and "截图" not in t):
                    break
                if t.startswith("---"):
                    i += 1
                    break
                i += 1
            continue

        if s.startswith("#### 📷"):
            # skip meta until image or next heading
            i += 1
            while i < len(lines):
                t = lines[i].strip()
                if t.startswith("![") or t.startswith("##") or t.startswith("###") or t.startswith("####"):
                    break
                if t.startswith("---"):
                    i += 1
                    break
                i += 1
            continue

        if in_code:
            if s.startswith("```"):
                add_code_block(doc, "\n".join(code_buf))
                code_buf = []
                in_code = False
            else:
                code_buf.append(raw)
            i += 1
            continue

        if s.startswith("```"):
            in_code = True
            code_buf = []
            i += 1
            continue

        if not s:
            i += 1
            continue

        if s == "---":
            i += 1
            continue

        if s.startswith("*全文完*") or s.startswith("*全文"):
            i += 1
            continue

        # image
        m_img = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", s)
        if m_img:
            alt, rel = m_img.group(1), m_img.group(2)
            path = (MD_PATH.parent / rel).resolve()
            caption = alt or path.stem
            add_image(doc, path, f"{fig_no}  {caption}")
            fig_no += 1
            i += 1
            continue

        # headings
        if s.startswith("## "):
            text = s[3:].strip()
            # strip anchor leftovers
            text = re.sub(r"\s*（含截图.*）\s*$", "", text)
            h = doc.add_heading(text, level=1)
            for run in h.runs:
                set_run_font(run, size=18, bold=True, color=BLUE)
            # decorative line under h1
            line_p = doc.add_paragraph()
            set_paragraph_spacing(line_p, before=0, after=8)
            add_horizontal_line(line_p)
            i += 1
            continue

        if s.startswith("### "):
            text = s[4:].strip()
            if "截图" in text and "约定" in text:
                i += 1
                continue
            h = doc.add_heading(text, level=2)
            for run in h.runs:
                set_run_font(run, size=14, bold=True, color=BLUE_ACCENT)
            i += 1
            continue

        if s.startswith("#### "):
            text = s[5:].strip()
            if text.startswith("📷"):
                i += 1
                continue
            h = doc.add_heading(text, level=3)
            for run in h.runs:
                set_run_font(run, size=12, bold=True, color=DARK)
            i += 1
            continue

        # table
        if s.startswith("|") and i + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[i + 1].strip()):
            headers, rows, ni = parse_table_block(lines, i)
            if headers and rows and not is_meta_table(headers, rows):
                # filter empty-ish
                add_table(doc, headers, rows)
            i = ni
            continue

        # blockquote
        if s.startswith(">"):
            text = s.lstrip("> ").strip()
            if text:
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=2, after=2)
                p.paragraph_format.left_indent = Cm(0.4)
                run = p.add_run(text.replace("**", ""))
                set_run_font(run, size=10, color=MUTED)
            i += 1
            continue

        # numbered list
        m_num = re.match(r"^(\d+)\.\s+(.*)$", s)
        if m_num:
            add_bullet(doc, m_num.group(2), numbered=True, index=int(m_num.group(1)))
            i += 1
            continue

        # bullet
        if s.startswith("- "):
            add_bullet(doc, s[2:])
            i += 1
            continue

        # skip html comments
        if s.startswith("<!--"):
            i += 1
            continue

        # normal paragraph
        add_body_paragraph(doc, s)
        i += 1


def main():
    md = MD_PATH.read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)
    setup_page(doc)
    add_cover(doc)
    add_toc_page(doc)
    convert_md(doc, md)

    # 结尾页
    doc.add_page_break()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(end, before=80)
    run = end.add_run("— 文档结束 —")
    set_run_font(run, size=12, color=MUTED)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("截图来源：系统实机页面（admin 演示环境）")
    set_run_font(run, size=10, color=MUTED)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Wrote {OUT_PATH}")
    print(f"Size: {OUT_PATH.stat().st_size / 1024 / 1024:.2f} MB")


if __name__ == "__main__":
    main()
